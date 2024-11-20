import pandas as pd
from docx import Document

# Data based on the document and filtering criteria
data = {
    "#": range(1, 9),  # Serial numbers
    "Code": ["54", "186535001", "186552002", "B00.1", "57920007", "54.2", "54.3", "9678009"],
    "Vocabulary": ["ICD9CM", "SNOMED", "SNOMED", "ICD10CM", "SNOMED", "ICD9CM", "ICD9CM", "SNOMED"],
    "Concept Name": [
        "Eczema herpeticum", "Eczema herpeticum", "Herpesviral vesicular dermatitis",
        "Herpesviral vesicular dermatitis", "Herpetic gingivostomatitis", "Herpetic gingivostomatitis",
        "Herpetic meningoencephalitis", "Herpetic meningoencephalitis"
    ],
    "RC": [142, 205, 4774, 4774, 1172, 486, 152, 152],
    "DRC": [142, 205, 4774, 4774, 1172, 486, 152, 152],
    "PC": [0, 128, 3457, 0, 972, 0, 0, 94],
    "DPC": [0, 128, 3457, 0, 972, 0, 0, 94],
    "Domain": ["Condition", "Condition", "Condition", "Condition", "Condition", "Condition", "Condition", "Condition"]
}

# Create a DataFrame
df = pd.DataFrame(data)

# Create a new Document
doc = Document()
doc.add_heading('Concept Table', 0)

# Add a table to the document
table = doc.add_table(rows=1, cols=9)

# Add the header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '#'
hdr_cells[1].text = 'Code'
hdr_cells[2].text = 'Vocabulary'
hdr_cells[3].text = 'Concept Name'
hdr_cells[4].text = 'RC'
hdr_cells[5].text = 'DRC'
hdr_cells[6].text = 'PC'
hdr_cells[7].text = 'DPC'
hdr_cells[8].text = 'Domain'

# Add data to the table
for i, row in df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['#'])
    row_cells[1].text = row['Code']
    row_cells[2].text = row['Vocabulary']
    row_cells[3].text = row['Concept Name']
    row_cells[4].text = str(row['RC'])
    row_cells[5].text = str(row['DRC'])
    row_cells[6].text = str(row['PC'])
    row_cells[7].text = str(row['DPC'])
    row_cells[8].text = row['Domain']

# Save the document
doc.save("Concept_Table.docx")

print("Document saved as 'Concept_Table.docx'")
