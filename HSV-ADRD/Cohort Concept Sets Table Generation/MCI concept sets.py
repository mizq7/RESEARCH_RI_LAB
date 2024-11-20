import pandas as pd
from docx import Document

# Data extracted based on the criteria
data = {
    "#": [1, 2, 3, 4, 5, 6],
    "Code": ["191519005", "386805003", "F06.71", "F06.70", "294.1", "331.83"],
    "Vocabulary": ["SNOMED", "SNOMED", "ICD10CM", "ICD10CM", "ICD9CM", "ICD9CM"],
    "Concept Name": [
        "Dementia associated with another disease",
        "Mild cognitive disorder",
        "Mild neurocognitive disorder with behavioral disturbance",
        "Mild neurocognitive disorder without behavioral disturbance",
        "Dementia in conditions classified elsewhere without behavioral disturbance",
        "Mild cognitive impairment so stated"
    ],
    "RC": [13994, 3091, 62, 328, 4273, 2896],
    "DRC": [32754, 3091, 62, 328, 4273, 2896],
    "PC": [6709, 2139, 0, 0, 0, 0],
    "DPC": [14940, 2139, 0, 0, 0, 0],
    "Domain": ["Condition", "Condition", "Condition", "Condition", "Condition", "Condition"]
}

# Create a DataFrame
df = pd.DataFrame(data)

# Saving the DataFrame as an Excel file
df.to_excel("MCI_Concept_Table.xlsx", index=False)

# Creating a Word document
doc = Document()
doc.add_heading('MCI Concept Table', 0)

# Adding a table to the document
table = doc.add_table(rows=1, cols=len(df.columns))

# Adding headers to the table
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(df.columns):
    hdr_cells[i].text = column_name

# Adding data to the table
for index, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# Saving the table to a Word document
doc.save("MCI_Concept_Table.docx")

# Displaying the DataFrame
print(df)

# Summarizing key information
unique_concepts = df['Concept Name'].nunique()
duplicated_concepts = df['Concept Name'].duplicated().sum()
unique_vocabularies = df['Vocabulary'].nunique()
vocabularies_list = df['Vocabulary'].unique().tolist()
unique_domains = df['Domain'].nunique()
pc_total = df['PC'].sum()

# Summary output
summary = {
    "Total Unique Concept Names": unique_concepts,
    "Duplicated Concept Names": duplicated_concepts,
    "Unique Vocabularies": unique_vocabularies,
    "Vocabulary Names": vocabularies_list,
    "Unique Domains": unique_domains,
    "Total PC Sum": pc_total
}

print("\nSummary of the Table:")
print(summary)
