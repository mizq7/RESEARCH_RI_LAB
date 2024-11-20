from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


def create_table_in_word(data):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add table
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['#', 'Code', 'Vocabulary', 'Concept Name', 'RC', 'DRC', 'PC', 'DPC', 'Domain']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Populate table
    serial_number = 1
    unique_concepts = set()
    unique_vocabularies = set()
    unique_domains = set()
    pc_sum = 0
    pc_sum_with_zero = 0

    for row in data:
        cells = table.add_row().cells
        cells[0].text = str(serial_number)
        cells[1].text = row[1]
        cells[2].text = row[8]
        cells[3].text = row[2]
        cells[4].text = row[3]
        cells[5].text = row[4]
        cells[6].text = row[5]
        cells[7].text = row[6]
        cells[8].text = row[7]

        serial_number += 1
        unique_concepts.add(row[2].lower())
        unique_vocabularies.add(row[8])
        unique_domains.add(row[7])

        pc = int(row[5].replace(',', '')) if row[5].replace(',', '').isdigit() else 0
        pc_sum += pc
        pc_sum_with_zero += pc

    # Add summary
    summary = (f"Total Unique Concept Names: {len(unique_concepts)}\n"
               f"Duplicated Concept Names: {serial_number - 1 - len(unique_concepts)}\n"
               f"Unique Vocabularies: {len(unique_vocabularies)}\n"
               f"Vocabulary Names: {', '.join(sorted(unique_vocabularies))}\n"
               f"Unique Domains: {len(unique_domains)} ({', '.join(sorted(unique_domains))})\n"
               f"Total PC Sum: {pc_sum:,}")
    doc.add_paragraph(summary)

    doc.add_paragraph(f"Note: The total PC sum including rows with PC=0 is: {pc_sum_with_zero:,}")

    # Add explanations
    doc.add_paragraph('RC (Record Count): total number of records')
    doc.add_paragraph('DRC (Distinct Record Count): number of unique records')
    doc.add_paragraph('PC (Patient Count): total number of patients')
    doc.add_paragraph('DPC (Distinct Patient Count): number of unique patients')

    # Save the document
    doc.save('all_cause_dementia_concepts.docx')
    print("Document 'all_cause_dementia_concepts.docx' has been created.")


# Your data
data = [
    ['37204495', '783161005', 'ABri amyloidosis', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['37204554', '783258000', 'ADan amyloidosis', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['45529867', '2901B', "ALZHEIMER'S DISEASE", '0', '0', '0', '0', 'Condition', 'OXMIS'],
    ['1002193', '93491-9', 'AMPAR2 IgG Ab [Presence] in Cerebral spinal fluid by Cell binding immunofluorescent assay',
     '0', '0', '0', '0', 'Measurement', 'LOINC'],
    ['4046095', '230301006', "Akinetic-rigid form of Huntington's disease", '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['37397762', '8.2351E+13', "Altered behavior due to Pick's disease", '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44784643', '9.7751E+13', "Altered behavior in Alzheimer's disease", '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44784620', '8.2361E+13', "Altered behavior in Huntington's dementia", '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['45618834', 'D000544', 'Alzheimer Disease', '0', '0', '0', '0', 'Condition', 'MeSH'],
    ['3296466', '26929004', 'Alzheimer dementia', '0', '0', '0', '0', 'Condition', 'Nebraska Lexicon'],
    ['45538101', 'G30', 'Alzheimer disease', '0', '0', '0', '0', 'Condition', 'ICD10'],
    ['37606320', 'G30', 'Alzheimer disease', '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['1411491', 'G30', 'Alzheimer disease', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['45616310', 'C536595', 'Alzheimer disease type 2', '0', '0', '0', '0', 'Condition', 'MeSH'],
    ['1411493', 'G30.000', 'Alzheimer disease with early onset', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['1411492', 'G30.0', 'Alzheimer disease with early onset', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['37083316', 'G30.0', 'Alzheimer disease with early onset', '0', '0', '0', '0', 'Condition', 'ICD10GM'],
    ['45566868', 'G30.0', 'Alzheimer disease with early onset', '0', '0', '0', '0', 'Condition', 'ICD10'],
    ['37610867', 'G30.0', 'Alzheimer disease with early onset', '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37083317', 'G30.1', 'Alzheimer disease with late onset', '0', '0', '0', '0', 'Condition', 'ICD10GM'],
    ['45552521', 'G30.1', 'Alzheimer disease with late onset', '0', '0', '0', '0', 'Condition', 'ICD10'],
    ['37607247', 'G30.1', 'Alzheimer disease with late onset', '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['1411495', 'G30.100', 'Alzheimer disease with late onset', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['1411494', 'G30.1', 'Alzheimer disease with late onset', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['37083319', 'G30.9', 'Alzheimer disease, unspecified', '0', '0', '0', '0', 'Condition', 'ICD10GM'],
    ['45591165', 'G30.9', 'Alzheimer disease, unspecified', '0', '0', '0', '0', 'Condition', 'ICD10'],
    ['1411501', 'G30.900', 'Alzheimer disease, unspecified', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['37611572', 'G30.9', 'Alzheimer disease, unspecified', '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['1411500', 'G30.9', 'Alzheimer disease, unspecified', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['45948811', '149019', "Alzheimer's Disease", '0', '0', '0', '0', 'Condition', 'CIEL'],
    ['45509950', 'F110.00', "Alzheimer's disease", '0', '0', '0', '0', 'Condition', 'Read'],
    ['378419', '26929004', "Alzheimer's disease", '23,124', '24,589', '5,888', '6,395', 'Condition', 'SNOMED'],
    ['36311271', 'LA22313-3', "Alzheimer's disease", '0', '0', '0', '0', 'Meas Value', 'LOINC'],
    ['1568293', 'G30', "Alzheimer's disease", '0', '0', '0', '0', 'Condition', 'ICD10CM'],
    ['1031954', 'LP74474-5', "Alzheimer's disease", '0', '0', '0', '0', 'Observation', 'LOINC'],
    ['44826537', '331', "Alzheimer's disease", '15,387', '15,387', '0', '0', 'Condition', 'ICD9CM'],
    ['3262766', '1.4201E+14', "Alzheimer's disease co-occurrent with delirium", '0', '0', '0', '0', 'Condition',
     'Nebraska Lexicon'],
    ['37395572', '1.4201E+14', "Alzheimer's disease co-occurrent with delirium", '0', '0', '0', '0', 'Condition',
     'SNOMED'],
    ['1411502', 'G30.901', "Alzheimer's disease dementia (machine translation)", '0', '0', '0', '0', 'Condition',
     'ICD10CN'],
    ['35207356', 'G30.0', "Alzheimer's disease with early onset", '311', '311', '0', '0', 'Condition', 'ICD10CM'],
    ['45513334', 'F110000', "Alzheimer's disease with early onset", '0', '0', '0', '0', 'Condition', 'Read'],
    ['35207357', 'G30.1', "Alzheimer's disease with late onset", '1,154', '1,154', '0', '0', 'Condition', 'ICD10CM'],
    ['45496654', 'F110100', "Alzheimer's disease with late onset", '0', '0', '0', '0', 'Condition', 'Read'],
    ['35207359', 'G30.9', "Alzheimer's disease, unspecified", '7,525', '7,525', '0', '0', 'Condition', 'ICD10CM'],
    ['35815705', '1010-10', "Alzheimer's disease/dementia", '0', '0', '0', '0', 'Observation', 'UK Biobank'],
    ['3170377', '2.4576E+14', 'Alzheimers continuum', '0', '0', '0', '0', 'Condition', 'Nebraska Lexicon'],
    ['42487047', 'G30.0', 'Alzheimer‚Äôs disease with early onset', '0', '0', '0', '0', 'Condition', 'KCD7'],
    ['42487048', 'G30.1', 'Alzheimer‚Äôs disease with late onset', '0', '0', '0', '0', 'Condition', 'KCD7'],
    ['42487050', 'G30.9', 'Alzheimer‚Äôs disease, unspecified', '0', '0', '0', '0', 'Condition', 'KCD7'],
    ['439147', '48167000', 'Amnesia', '33,654', '34,749', '15,812', '16,663', 'Condition', 'SNOMED'],
    ['4092086', '247607004', 'Amnesia for day to day facts', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['4083456', '247611005', 'Amnesia for important personal information', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['4171718', '42176003', 'Amnesia for recent events', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['4264146', '6149008', 'Amnesia for remote events', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44821813', '294', 'Amnestic disorder in conditions classified elsewhere', '219', '219', '0', '0', 'Condition',
     'ICD9CM'],
    ['3654469', '836301008', 'Amnestic mild cognitive disorder', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['36660072', '94354-8', 'Amphiphysin IgG Ab [Titer] in Cerebral spinal fluid by Immunofluorescence', '0', '0', '0',
     '0', 'Measurement', 'LOINC'],
    ['3654598', '838276009', 'Amyotrophic lateral sclerosis, parkinsonism, dementia complex', '0', '0', '0', '0',
     'Condition', 'SNOMED'],
    ['4229448', '88822006', 'Anterograde amnesia', '431', '431', '369', '369', 'Condition', 'SNOMED'],
    ['4046087', '230281007', 'Argyrophilic grain disease', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['376094', '191464005', 'Arteriosclerotic dementia with delirium', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['374326', '191466007', 'Arteriosclerotic dementia with depression', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['3348091', '191465006', 'Arteriosclerotic dementia with paranoia', '0', '0', '0', '0', 'Condition',
     'Nebraska Lexicon'],
    ['4100252', '191465006', 'Arteriosclerotic dementia with paranoia', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['45490025', 'E004200', 'Arteriosclerotic dementia with paranoia', '0', '0', '0', '0', 'Condition', 'Read'],
    ['1011812', 'LP156130-9', 'Ascertain Dementia 8', '0', '0', '0', '0', 'Observation', 'LOINC'],
    ['42869700', '71945-0', 'Ascertain Dementia 8 [AD8]', '0', '0', '0', '0', 'Observation', 'LOINC'],
    ['608051', '1156789004', 'Autosomal dominant Alzheimer disease due to mutation of amyloid precursor protein', '0',
     '0', '0', '0', 'Condition', 'SNOMED'],
    ['603149', '1156800008', 'Autosomal dominant Alzheimer disease due to mutation of presenilin 1', '0', '0', '0', '0',
     'Condition', 'SNOMED'],
    ['608060', '1156798001', 'Autosomal dominant Alzheimer disease due to mutation of presenilin 2', '0', '0', '0', '0',
     'Condition', 'SNOMED'],
    ['37117145', '1.6219E+16', 'Behavioral disturbance co-occurrent and due to late onset Alzheimer dementia', '0', '0',
     '0', '0', 'Condition', 'SNOMED'],
    ['37399020', '716994006', 'Behavioral variant of frontotemporal dementia', '0', '0', '0', '0', 'Condition',
     'SNOMED'],
    ['45771254', '702393003', 'CHMP2B-related frontotemporal dementia', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['3037012', '35385-4', 'CV2 IgG Ab [Titer] in Cerebral spinal fluid', '0', '0', '0', '0', 'Measurement', 'LOINC'],
    ['4092747', '279982005', 'Cerebral degeneration presenting primarily with dementia', '0', '27,496', '0', '7,393',
     'Condition', 'SNOMED'],
    ['36660386', '94286-2',
     'Contactin-associated protein 2 IgG Ab [Presence] in Cerebral spinal fluid by Cell binding immunofluorescent assay',
     '0', '0', '0', '0', 'Measurement', 'LOINC'],
    ['3654434', '833326008', 'Cortical vascular dementia', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['37111242', '725898002', 'Delirium co-occurrent with dementia', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44782726', '1.4199E+14', "Delusions in Alzheimer's disease", '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['4182210', '52448006', 'Dementia', '30,560', '110,249', '11,423', '41,063', 'Condition', 'SNOMED'],
    ['3325472', '4817008', 'Dementia Alzheimers type, late onset with delirium', '0', '0', '0', '0', 'Condition',
     'Nebraska Lexicon'],
    ['4228133', '421529006', 'Dementia associated with AIDS', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['4314734', '425390006', "Dementia associated with Parkinson's Disease", '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['378726', '281004', 'Dementia associated with alcoholism', '487', '487', '237', '237', 'Condition', 'SNOMED'],
    ['374888', '191519005', 'Dementia associated with another disease', '13,994', '32,754', '6,709', '14,940',
     'Condition', 'SNOMED'],
    ['44784607', '698781002', 'Dementia associated with cerebral anoxia', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44784472', '698624003', 'Dementia associated with cerebral lipidosis', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44784474', '698626001', 'Dementia associated with multiple sclerosis', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44784559', '698725008', 'Dementia associated with neurosyphilis', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44784473', '698625002', 'Dementia associated with normal pressure hydrocephalus', '0', '0', '0', '0', 'Condition',
     'SNOMED'],
    ['44784560', '698726009', 'Dementia associated with viral encephalitis', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['36659857', '94707-7', 'Dementia autoimmune Ab panel - Cerebral spinal fluid', '0', '0', '0', '0', 'Measurement',
     'LOINC'],
    ['36660881', 'LP419388-6', 'Dementia autoimmune Ab panel | Cerebral spinal fluid | Serology Panels', '0', '0', '0',
     '0', 'Measurement', 'LOINC'],
    ['37116464', '733184002', 'Dementia caused by heavy metal exposure', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['36717598', '722978000', 'Dementia caused by toxin', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['37311999', '788898005', 'Dementia caused by volatile inhalant', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['36716795', '722977005', 'Dementia co-occurrent and due to neurocysticercosis', '0', '0', '0', '0', 'Condition',
     'SNOMED'],
    ['37017549', '713844000', 'Dementia co-occurrent with human immunodeficiency virus infection', '0', '0', '0', '0',
     'Condition', 'SNOMED'],
    ['4180284', '429458009', 'Dementia due to Creutzfeldt Jakob disease', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['40483103', '442344002', 'Dementia due to Huntington chorea', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44782422', '1.0142E+14', "Dementia due to Parkinson's disease", '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44782710', '2.1921E+13', "Dementia due to Pick's disease", '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['43020422', '1.3012E+14', "Dementia due to Rett's syndrome", '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['42538610', '762351006', 'Dementia due to and following injury of head', '0', '0', '0', '0', 'Condition',
     'SNOMED'],
    ['3654920', '840464007', 'Dementia due to carbon monoxide poisoning', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['36716797', '722980006', 'Dementia due to chromosomal anomaly', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['37116466', '733191004', 'Dementia due to chronic subdural hematoma', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['606974', '1148924004', 'Dementia due to deficiency of folic acid', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['37110513', '724776007', 'Dementia due to disorder of central nervous system', '0', '0', '0', '0', 'Condition',
     'SNOMED'],
    ['37116467', '733192006', 'Dementia due to herpes encephalitis', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['37119154', '724777003', 'Dementia due to infectious disease', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['3654921', '840465008', 'Dementia due to iron deficiency', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['36716796', '722979008', 'Dementia due to metabolic abnormality', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['44782559', '8.2371E+13', 'Dementia due to multiple sclerosis with altered behavior', '0', '0', '0', '0',
     'Condition', 'SNOMED'],
    ['37311998', '788899002', 'Dementia due to pellagra', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['37116465', '733190003', 'Dementia due to primary malignant neoplasm of brain', '0', '0', '0', '0', 'Condition',
     'SNOMED'],
    ['42538609', '762350007', 'Dementia due to prion disease', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['42535731', '733185001', 'Dementia following injury caused by exposure to ionizing radiation', '0', '0', '0', '0',
     'Condition', 'SNOMED'],
    ['37082547', 'F00.0', 'Dementia in Alzheimer disease with early onset', '0', '0', '0', '0', 'Condition', 'ICD10GM'],
    ['1410202', 'F00.0', 'Dementia in Alzheimer disease with early onset', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['37611340', 'F00.0', 'Dementia in Alzheimer disease with early onset', '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['45591070', 'F00.0', 'Dementia in Alzheimer disease with early onset', '0', '0', '0', '0', 'Condition', 'ICD10'],
    ['37615601', 'F00.1', 'Dementia in Alzheimer disease with late onset', '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37082548', 'F00.1', 'Dementia in Alzheimer disease with late onset', '0', '0', '0', '0', 'Condition', 'ICD10GM'],
    ['1410203', 'F00.1', 'Dementia in Alzheimer disease with late onset', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['45755371', 'F00.1', 'Dementia in Alzheimer disease with late onset', '0', '0', '0', '0', 'Condition', 'ICD10'],
    ['37610391', 'F00.2', 'Dementia in Alzheimer disease, atypical or mixed type', '0', '0', '0', '0', 'Condition',
     'CIM10'],
    ['37082549', 'F00.2', 'Dementia in Alzheimer disease, atypical or mixed type', '0', '0', '0', '0', 'Condition',
     'ICD10GM'],
    ['1410204', 'F00.2', 'Dementia in Alzheimer disease, atypical or mixed type', '0', '0', '0', '0', 'Condition',
     'ICD10CN'],
    ['45557146', 'F00.2', 'Dementia in Alzheimer disease, atypical or mixed type', '0', '0', '0', '0', 'Condition',
     'ICD10'],
    ['37610454', 'F00.9', 'Dementia in Alzheimer disease, unspecified', '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37082550', 'F00.9', 'Dementia in Alzheimer disease, unspecified', '0', '0', '0', '0', 'Condition', 'ICD10GM'],
    ['1410205', 'F00.9', 'Dementia in Alzheimer disease, unspecified', '0', '0', '0', '0', 'Condition', 'ICD10CN'],
    ['45561994', 'F00.9', 'Dementia in Alzheimer disease, unspecified', '0', '0', '0', '0', 'Condition', 'ICD10'],
    ['45950909', '152127', "Dementia in Alzheimer's Disease with Early Onset", '0', '0', '0', '0', 'Condition', 'CIEL'],
    ['3319436', '6475002', "Dementia in Alzheimer's disease with early onset", '0', '0', '0', '0', 'Condition',
     'Nebraska Lexicon'],
    ['3216550', '66108005', "Dementia in Alzheimer's disease with late onset", '0', '0', '0', '0', 'Condition',
     'Nebraska Lexicon'],
    ['3313390', '416975007', "Dementia in Alzheimer's disease with late onset", '0', '0', '0', '0', 'Condition',
     'Nebraska Lexicon'],
    ['42486488', 'F00.9', 'Dementia in Alzheimer¬¥s disease, unspecified(G30.9‚Ä†)', '0', '0', '0', '0', 'Condition',
     'KCD7'],
    ['42486485', 'F00.0', 'Dementia in Alzheimer‚Äôs disease with early onset(G30.0‚Ä†)', '0', '0', '0', '0',
     'Condition', 'KCD7'],
    ['42486486', 'F00.1', 'Dementia in Alzheimer‚Äôs disease with late onset(G30.1‚Ä†)', '0', '0', '0', '0',
     'Condition', 'KCD7'],
    ['42486487', 'F00.2', 'Dementia in Alzheimer‚Äôs disease, atypical or mixed type(G30.8‚Ä†)', '0', '0', '0', '0',
     'Condition', 'KCD7'],
    ['44782935', '698949001', 'Dementia in remission', '0', '0', '0', '0', 'Condition', 'SNOMED'],
    ['37605285', 'F00.24', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mixed", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604711', 'F00.240', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mixed, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604712', 'F00.241', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mixed, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604713', 'F00.242', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mixed, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605282', 'F00.21', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly delusional", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604705', 'F00.210', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly delusional, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604706', 'F00.211', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly delusional, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604707', 'F00.212', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly delusional, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605284', 'F00.23', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly depressive", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605177', 'F00.230', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly depressive, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605178', 'F00.231', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly depressive, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605179', 'F00.232', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly depressive, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605283', 'F00.22', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly hallucinatory", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604756', 'F00.220', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly hallucinatory, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604757', 'F00.221', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly hallucinatory, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604758', 'F00.222', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, with other symptoms, mostly hallucinatory, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605281', 'F00.20', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, without additional symptoms", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604838', 'F00.200', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, without additional symptoms, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604839', 'F00.201', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, without additional symptoms, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604840', 'F00.202', "Dementia of Alzheimer's disease, atypical or mixed form G30.8, without additional symptoms, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605271', 'F00.00', "Dementia of Alzheimer's disease, early onset G30.0, no additional symptoms", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604915', 'F00.000', "Dementia of Alzheimer's disease, early onset G30.0, no additional symptoms, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604917', 'F00.001', "Dementia of Alzheimer's disease, early onset G30.0, no additional symptoms, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604918', 'F00.002', "Dementia of Alzheimer's disease, early onset G30.0, no additional symptoms, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605275', 'F00.04', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mixed", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604774', 'F00.040', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mixed, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604776', 'F00.041', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mixed, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604777', 'F00.042', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mixed, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605272', 'F00.01', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly delusional", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604796', 'F00.010', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly delusional, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604797', 'F00.011', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly delusional, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604798', 'F00.012', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly delusional, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605274', 'F00.03', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly depressive", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604831', 'F00.030', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly depressive, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604832', 'F00.031', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly depressive, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604833', 'F00.032', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly depressive, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605273', 'F00.02', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly hallucinatory", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604802', 'F00.020', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly hallucinatory, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604803', 'F00.021', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly hallucinatory, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604804', 'F00.022', "Dementia of Alzheimer's disease, early onset G30.0, with other symptoms, mostly hallucinatory, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605280', 'F00.14', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mixed", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604750', 'F00.140', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mixed, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604751', 'F00.141', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mixed, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604752', 'F00.142', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mixed, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605277', 'F00.11', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mostly delusional", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605165', 'F00.110', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mostly delusional, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605166', 'F00.111', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mostly delusional, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605167', 'F00.112', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mostly delusional, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605279', 'F00.13', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mostly depressive", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604919', 'F00.130', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mostly depressive, mild", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604920', 'F00.131', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mostly depressive, moderate", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37604921', 'F00.132', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mostly depressive, severe", '0', '0', '0', '0', 'Condition', 'CIM10'],
    ['37605278', 'F00.12', "Dementia of Alzheimer's disease, late onset G30.1, with other symptoms, mostly hallucinatory", '0', '0', '0', '0', 'Condition', 'CIM10'],
]

create_table_in_word(data)