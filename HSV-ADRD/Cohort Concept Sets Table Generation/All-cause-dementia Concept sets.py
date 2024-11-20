import pandas as pd
from docx import Document
from docx.shared import Inches

# Create DataFrame from the provided data
data = [
    [37204495, '783161005', 'ABri amyloidosis', 0, 0, 0, 0, 'Condition', 'SNOMED'],
    [37204554, '783258000', 'ADan amyloidosis', 0, 0, 0, 0, 'Condition', 'SNOMED'],
    [45529867, '2901B', "ALZHEIMER'S DISEASE", 0, 0, 0, 0, 'Condition', 'OXMIS'],
    # ... (insert all other data rows here)
    [37605315, 'F06.71', 'mild cognitive impairment associated with a physical disorder', 0, 0, 0, 0, 'Condition', 'CIM10']
]

columns = ['Id', 'Code', 'Name', 'RC', 'DRC', 'PC', 'DPC', 'Domain', 'Vocabulary']
df = pd.DataFrame(data, columns=columns)

# Convert numeric columns to appropriate types
numeric_columns = ['RC', 'DRC', 'PC', 'DPC']
df[numeric_columns] = df[numeric_columns].apply(pd.to_numeric, errors='coerce')

# Generate summary
unique_concepts = df['Name'].nunique()
total_concepts = len(df)
duplicated_concepts = total_concepts - unique_concepts
unique_vocabularies = df['Vocabulary'].nunique()
vocabulary_names = df['Vocabulary'].unique().tolist()
unique_domains = df['Domain'].nunique()
total_pc_sum = df['PC'].sum()

# Create a new Word document
doc = Document()

# Add a title
doc.add_heading('All-Cause Dementia Concept Sets Summary', 0)

# Add a table with the summary information
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'

# Populate the table
data = [
    ("Total Unique Concept Names", str(unique_concepts)),
    ("Duplicated Concept Names", str(duplicated_concepts)),
    ("Unique Vocabularies", str(unique_vocabularies)),
    ("Vocabulary Names", ", ".join(vocabulary_names)),
    ("Unique Domains", f"{unique_domains} ({', '.join(df['Domain'].unique())})"),
    ("Total PC Sum", str(total_pc_sum))
]

for i, (key, value) in enumerate(data):
    table.cell(i, 0).text = key
    table.cell(i, 1).text = value

# Add explanations
doc.add_paragraph("\nMetric Explanations:")
doc.add_paragraph("RC (Record Count): total number of records")
doc.add_paragraph("DRC (Distinct Record Count): number of unique records")
doc.add_paragraph("PC (Patient Count): total number of patients")
doc.add_paragraph("DPC (Distinct Patient Count): number of unique patients")

# Save the document
doc.save('all_cause_dementia_summary.docx')

print("Summary has been saved to 'all_cause_dementia_summary.docx'")